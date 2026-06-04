import re
import docx

def parse_docx(file_stream):
    try:
        doc = docx.Document(file_stream)
        paragraphs = [para.text for para in doc.paragraphs]
        text = "\n".join(paragraphs)
        
        # Extract text from tables if present
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text for cell in row.cells if cell.text]
                table_texts.append(" ".join(row_texts))
        if table_texts:
            text += "\n" + "\n".join(table_texts)
            
        if not text:
            return ""
        clean_text = re.sub(r'\s+', ' ', text).strip()
        return clean_text
    except Exception as e:
        raise RuntimeError(f"DOCX parsing error: {str(e)}") from e
